"""
Parsing layer for the LegalRAG ingestion pipeline.

Responsibilities
-----------------
1. Load the raw law corpus (`corpus_law_pub.json`, the file referenced by
   `docs/submission_example.json` as the source of `aid`s) and normalize it
   into a flat list of `RawArticle` records.
2. Optionally parse supplementary law documents supplied as .docx/.pdf, for
   teams that extend the corpus beyond what the organizers provide.
3. Split the body of each article into its Chương > Mục > Điều > Khoản > Điểm
   structure using rule-based regex (never token-based cutting), because a
   fixed-size split can sever a clause from the "trừ trường hợp..." exception
   that governs its meaning.

This module intentionally does NOT do embedding/indexing — see
`backend/indexing/`. It only produces clean, structured text ready to be
chunked (`backend/ingestion/chunker.py`).
"""
from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Optional

# --- Regex patterns for Vietnamese legal structure markers ----------------
# Matches "Chương I", "Chương II", ...
RE_CHUONG = re.compile(r"^\s*Chương\s+[IVXLCDM\d]+\b.*$", re.MULTILINE)
# Matches "Mục 1", "Mục 2", ...
RE_MUC = re.compile(r"^\s*Mục\s+\d+\b.*$", re.MULTILINE)
# Matches "Điều 12.", "Điều 12:" at line start
RE_DIEU = re.compile(r"^\s*Điều\s+(\d+)\s*[.:]?\s*(.*)$", re.MULTILINE)
# Matches "1.", "2." at line start (Khoản), only within an Điều block
RE_KHOAN = re.compile(r"^\s*(\d+)\s*\.\s+")
# Matches "a)", "b)" at line start (Điểm)
RE_DIEM = re.compile(r"^\s*([a-zđ])\s*\)\s+", re.IGNORECASE)


@dataclass
class RawArticle:
    law_id: str
    aid: int
    title: str = ""
    chuong: Optional[str] = None
    muc: Optional[str] = None
    body: str = ""  # full article text, Khoản/Điểm markers preserved


@dataclass
class RawLawDocument:
    law_id: str
    doc_type: Optional[str] = None
    issuing_body: Optional[str] = None
    issue_date: Optional[str] = None
    effective_date: Optional[str] = None
    expiry_date: Optional[str] = None
    articles: list[RawArticle] = field(default_factory=list)


def load_law_corpus(path: str | Path) -> list[RawLawDocument]:
    """Load `corpus_law_pub.json`.

    Expected (organizer-provided) shape, per docs/submission_example.json
    (each article has a `law_id` + integer `aid`):

    ```json
    [
      {
        "law_id": "47/2010/QH12",
        "doc_type": "Luật",
        "issuing_body": "Quốc hội",
        "issue_date": "2010-06-17",
        "effective_date": "2011-01-01",
        "articles": [
          {"aid": 270, "title": "Điều 270. ...", "text": "1. ...\\na) ...\\n2. ..."}
        ]
      }
    ]
    ```

    The loader is defensive about minor schema variations (e.g. `content`
    instead of `text`, `id` instead of `aid`) since the organizers' exact
    field names can shift between releases.
    """
    path = Path(path)
    raw = json.loads(path.read_text(encoding="utf-8"))
    if isinstance(raw, dict):
        raw = raw.get("laws") or raw.get("data") or [raw]

    docs: list[RawLawDocument] = []
    for law in raw:
        law_id = str(law.get("law_id") or law.get("id"))
        doc = RawLawDocument(
            law_id=law_id,
            doc_type=law.get("doc_type") or law.get("type"),
            issuing_body=law.get("issuing_body"),
            issue_date=law.get("issue_date"),
            effective_date=law.get("effective_date"),
            expiry_date=law.get("expiry_date"),
        )
        for art in law.get("articles", []):
            aid = art.get("aid")
            if aid is None:
                aid = art.get("id")
            text = art.get("text") or art.get("content") or ""
            doc.articles.append(
                RawArticle(
                    law_id=law_id,
                    aid=int(aid),
                    title=art.get("title", ""),
                    chuong=art.get("chuong") or art.get("chapter"),
                    muc=art.get("muc") or art.get("section"),
                    body=text.strip(),
                )
            )
        docs.append(doc)
    return docs


def load_test_set(path: str | Path) -> list[dict]:
    """Load `ALQAC2026_public_test.json` (list of {case_id, case_query, ...})."""
    path = Path(path)
    data = json.loads(path.read_text(encoding="utf-8"))
    if isinstance(data, dict):
        data = data.get("data") or data.get("cases") or []
    return data


# ---------------------------------------------------------------------------
# Structural splitting (Khoản / Điểm) — used by chunker.py
# ---------------------------------------------------------------------------
@dataclass
class KhoanSplit:
    khoan_no: Optional[str]
    diem_no: Optional[str]
    text: str


def split_article_into_khoan_diem(body: str) -> list[KhoanSplit]:
    """Split an article body into its Khoản (and nested Điểm) units.

    Falls back to returning the whole body as a single unit if no Khoản
    markers are detected (some articles are a single unstructured paragraph).
    """
    lines = [l for l in body.splitlines() if l.strip()]
    if not lines:
        return []

    units: list[KhoanSplit] = []
    current_khoan: Optional[str] = None
    current_lines: list[str] = []
    current_diem: Optional[str] = None

    def flush():
        if current_lines:
            units.append(
                KhoanSplit(
                    khoan_no=current_khoan,
                    diem_no=current_diem,
                    text=" ".join(current_lines).strip(),
                )
            )

    for line in lines:
        khoan_match = RE_KHOAN.match(line)
        diem_match = RE_DIEM.match(line)
        if khoan_match:
            flush()
            current_khoan = khoan_match.group(1)
            current_diem = None
            current_lines = [line[khoan_match.end():].strip()]
        elif diem_match:
            flush()
            current_diem = diem_match.group(1)
            current_lines = [line[diem_match.end():].strip()]
        else:
            current_lines.append(line.strip())
    flush()

    if not units:
        units = [KhoanSplit(khoan_no=None, diem_no=None, text=body.strip())]
    return units


def parse_docx_supplement(path: str | Path) -> str:
    """Best-effort text extraction from a supplementary .docx law file.

    Requires `python-docx`. Only used when teams extend the corpus with
    documents not already provided as JSON by the organizers.
    """
    try:
        import docx  # python-docx
    except ImportError as e:
        raise ImportError("python-docx is required to parse .docx files: pip install python-docx") from e

    d = docx.Document(str(path))
    return "\n".join(p.text for p in d.paragraphs if p.text.strip())


def parse_pdf_supplement(path: str | Path) -> str:
    """Best-effort text extraction from a supplementary .pdf law file."""
    try:
        import pdfplumber
    except ImportError as e:
        raise ImportError("pdfplumber is required to parse .pdf files: pip install pdfplumber") from e

    text_parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            text_parts.append(t)
    return "\n".join(text_parts)
